# pages/5_Report.py
# FINAL REPORT COMPILER — WIRE-FRAME ALIGNED (TEC 7.1 COMPLIANT)
# UI REFINED + SECTION STATUS (NO LOGIC CHANGES)

import streamlit as st
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches
import tempfile
import matplotlib.pyplot as plt

from utils.viz_utils import (
    plot_disparity_in_performance,
    plot_group_error_panel,
)



# ==================================================
# PAGE CONFIG
# ==================================================
st.set_page_config(layout="wide")
st.title("Final Fairness Evaluation Report")

st.markdown(
    """
    <style>
    /* Import font */
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;600;700;800&display=swap');
    
    /* Base typography */
    .stApp {
        font-family: 'Plus Jakarta Sans', sans-serif !important;
    }
    
    /* Gradient headings */
    h1, h2, h3 {
        background: linear-gradient(135deg, #1e40af 0%, #3b82f6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800 !important;
        letter-spacing: -0.02em;
    }
    
    /* Navigation sidebar - boxed style */
    [data-testid="stSidebarNav"] {
        background-color: rgba(255, 255, 255, 0.05) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-radius: 12px !important;
        padding: 1rem !important;
        margin: 1rem 0.5rem !important;
    }
    
    [data-testid="stSidebarNav"] a {
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-radius: 8px !important;
        padding: 0.75rem 1rem !important;
        margin-bottom: 0.5rem !important;
        transition: all 0.2s ease !important;
    }
    
    [data-testid="stSidebarNav"] a:hover {
        border-color: #60a5fa !important;
        background-color: rgba(96, 165, 250, 0.1) !important;
        transform: translateX(4px);
    }
    
    [data-testid="stSidebarNav"] a[aria-current="page"] {
        background: linear-gradient(135deg, #2563eb 0%, #0ea5e9 100%) !important;
        border-color: transparent !important;
    }
    
    /* Content cards - smaller padding */
    div[data-testid="stVerticalBlock"] > div[data-testid="stVerticalBlockBorderWrapper"] {
        border-radius: 16px !important;
        padding: 1.5rem !important;
        margin-bottom: 1rem !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
    }
    
    /* Radio button containers */
    div[data-testid="stRadio"] > div {
        background-color: rgba(255, 255, 255, 0.03) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-radius: 12px !important;
        padding: 1rem !important;
    }
    
    /* Primary buttons */
    button[kind="primary"] {
        background: linear-gradient(135deg, #2563eb 0%, #0ea5e9 100%) !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.6rem 1.5rem !important;
        font-weight: 700 !important;
    }
    
    /* File uploader */
    div[data-testid="stFileUploader"] {
        border-radius: 12px !important;
        padding: 1.5rem !important;
    }
    
    /* Metrics */
    div[data-testid="stMetric"] {
        border-radius: 12px !important;
        padding: 1rem !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
    }
    
    /* Logo positioning */
    .nishpaksh-logo {
        position: fixed;
        top: 16px;
        left: 16px;
        width: 120px;
        z-index: 9999;
    }
    
    [data-testid="stSidebarNav"] {
        margin-top: 120px;
    }

    /* ===============================
   Boxed & symmetric tabs
   =============================== */
   div[data-testid="stTabs"] {
    padding: 0.25rem 0.25rem 0.75rem 0.25rem;
    }

    div[data-testid="stTabs"] button {
    border: 1px solid rgba(255, 255, 255, 0.18) !important;
    border-radius: 10px !important;
    padding: 0.6rem 1.1rem !important;
    margin-right: 0.5rem !important;
    font-weight: 600 !important;
    background-color: rgba(255, 255, 255, 0.04) !important;
    transition: all 0.2s ease-in-out;
    }

    /* Active tab */
    div[data-testid="stTabs"] button[aria-selected="true"] {
    background: linear-gradient(135deg, #2563eb 0%, #0ea5e9 100%) !important;
    color: white !important;
    border-color: transparent !important;
    }

    /* Hover */
    div[data-testid="stTabs"] button:hover {
    border-color: #60a5fa !important;
    background-color: rgba(96, 165, 250, 0.12) !important;
    }

    
    /* Scrollbar */
    ::-webkit-scrollbar {
        width: 8px;
        height: 8px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(135deg, #2563eb 0%, #0ea5e9 100%);
        border-radius: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ==================================================
# IMPORTS REQUIRED FOR THIS PAGE LOGIC
# ==================================================
import numpy as np
import tempfile
import matplotlib.pyplot as plt
from pathlib import Path
from docx import Document
from docx.shared import Inches

from utils.viz_utils import (
    plot_disparity_in_performance,
    plot_group_error_panel,
)

# ==================================================
# DOCX HELPER FUNCTIONS (NO EXTERNAL DEPENDENCIES)
# ==================================================
def replace_text(doc: Document, placeholder: str, value: str) -> bool:
    found = False

    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, str(value))
            found = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, str(value))
                        found = True

    return found


def insert_plot(doc: Document, placeholder: str, fig) -> bool:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(tmp.name, dpi=800, bbox_inches="tight")
    plt.close(fig)

    found = False

    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, "")
            p.add_run().add_picture(tmp.name, width=Inches(6.5))
            found = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        p.add_run().add_picture(tmp.name, width=Inches(6.5))
                        found = True

    return found

def build_survey_qa_section_from_submission(survey_outputs):
    """
    Build Section 3.3 text from stored survey submission.
    Uses section names + question IDs + responses.
    No dependency on survey schema.
    """

    submission = survey_outputs.get("raw_submission", {})
    answers = submission.get("answers", {})

    if not answers:
        return "No survey responses were recorded."

    lines = []
    lines.append("Fairness and Governance Questionnaire Responses:\n")

    for section, qmap in answers.items():
        lines.append(f"{section.replace('_', ' ').title()}:")
        for q_id, response in qmap.items():
            lines.append(f"  - {q_id}: {response}")
        lines.append("")

    return "\n".join(lines)


# ==================================================
# SESSION STATE — SINGLE SOURCE OF TRUTH
# ==================================================
REQUIRED_KEYS = [
    "survey_outputs",
    "preproc",
    "inference",
    "results",
    "uploaded_data",
]

for k in REQUIRED_KEYS:
    if k not in st.session_state:
        st.error(f"Missing required step: {k}")
        st.stop()

survey = st.session_state["survey_outputs"]
preproc = st.session_state["preproc"]
inference = st.session_state["inference"]
results = st.session_state["results"]
df = st.session_state["uploaded_data"]

# ==================================================
# DEFENSIVE INITIALIZATION (ABSOLUTE)
# ==================================================
survey.setdefault("risk_bucket", "Medium")
survey.setdefault("auditor_access", "")
survey.setdefault("testing_type", "")
survey.setdefault("dependencies", "")
survey.setdefault("limitations", "")
survey.setdefault("questionnaire_summary", "")
survey.setdefault("privileged_groups", "")
survey.setdefault("favourable_outcome", "")
survey.setdefault("protected_attr_rationale", "")
survey.setdefault("metric_rationale", "")
survey.setdefault("threshold_rationale", "")
survey.setdefault("risk_outcome", "")
survey.setdefault("certification_context", "")

preproc.setdefault("user_narratives", {})
preproc["user_narratives"].setdefault("P14", "")
preproc.setdefault("pipeline_description", "")

# ==================================================
# HARD RESULTS CONTRACT (NEW FAIRNESS MATH)
# ==================================================
for k in ["selected_model", "FS_system", "BI_per_attribute", "verdict"]:
    if k not in results:
        st.error(f"Results object missing key: {k}")
        st.stop()

protected_attrs = list(results["BI_per_attribute"].keys())
if not protected_attrs:
    st.error("No protected attributes found.")
    st.stop()

# ==================================================
# PAGE TABS
# ==================================================
tab_summary, tab_metrics, tab_narrative, tab_report = st.tabs(
    ["Summary", "Fairness Metrics", "Narrative Inputs", "Generate Report"]
)

# ==================================================
# ==================================================
# TAB 1 — SUMMARY
# ==================================================
with tab_summary:

    # ==================================================
    # ORGANIZATIONAL CONTEXT
    # ==================================================
    st.markdown("## Organizational Context")

    with st.container(border=True):

        survey["organization_name"] = st.text_input(
            "Organization / Entity",
            survey.get("organization_name", ""),
            help=(
                "Legal entity or organization responsible for the development, "
                "deployment, and governance of this AI system."
            ),
        )

        survey["developing_department"] = st.text_input(
            "Department / Business Unit",
            survey.get("developing_department", ""),
            help=(
                "Internal department, team, or business unit that owns or "
                "operates the AI system."
            ),
        )

    # ==================================================
    # AI APPLICATION OVERVIEW
    # ==================================================
    st.markdown("## AI Application Overview")

    with st.container(border=True):

        survey["ai_application_overview"] = st.text_area(
            "High-level description of the AI application",
            survey.get("ai_application_overview", ""),
            height=100,
            help=(
                "Describe the purpose, scope, target users, and decision context "
                "of the AI system."
            ),
        )

    # ==================================================
    # FAIRNESS CONTEXT
    # ==================================================
    st.markdown("## Fairness Context")

    with st.container(border=True):

        survey["privileged_groups"] = st.text_area(
            "Privileged / Reference Groups",
            survey.get("privileged_groups", ""),
            height=80,
            help=(
                "Reference or baseline groups used for fairness comparison "
                "(e.g., majority group, control cohort)."
            ),
        )

        survey["favourable_outcome"] = st.text_area(
            "Definition of Favourable Outcome",
            survey.get("favourable_outcome", ""),
            height=80,
            help=(
                "Define what constitutes a positive or beneficial outcome "
                "in this AI system."
            ),
        )

    # ==================================================
    # AUDIT & GOVERNANCE CONTEXT  
    # ==================================================
    st.markdown("## Audit and Governance Context")

    with st.container(border=True):

        survey["risk_bucket"] = st.selectbox(
            "Audit Type / Risk Classification",
            ["Low", "Medium", "High"],
            index=["Low", "Medium", "High"].index(
                survey.get("risk_bucket", "Medium")
            ),
            help=(
                "Overall audit classification based on regulatory exposure, "
                "deployment criticality, and potential harm."
            ),
        )

        survey["auditor_access"] = st.text_area(
            "Auditor Access to Data",
            survey.get("auditor_access", ""),
            height=80,
            help=(
                "Describe the level of access provided to internal or external "
                "auditors (e.g., full data, anonymized samples, metrics only)."
            ),
        )

        survey["testing_type"] = st.text_area(
            "Testing Type",
            survey.get("testing_type", ""),
            height=80,
            help=(
                "Specify the evaluation methodology used "
                "(e.g., pre-deployment testing, post-deployment monitoring, "
                "periodic audit, shadow testing)."
            ),
        )

    # ==================================================
    # SYSTEM DEPENDENCIES & LIMITATIONS
    # ==================================================
    st.markdown("## System Dependencies and Limitations")

    with st.container(border=True):

        survey["dependencies"] = st.text_area(
            "Dependencies",
            survey.get("dependencies", ""),
            height=80,
            help=(
                "Key technical, data, infrastructure, or organizational "
                "dependencies required for correct operation."
            ),
        )

        survey["limitations"] = st.text_area(
            "Limitations",
            survey.get("limitations", ""),
            height=80,
            help=(
                "Known constraints, assumptions, data gaps, or evaluation "
                "limitations relevant to fairness assessment."
            ),
        )
# ==================================================
# TAB 2 — FAIRNESS METRICS (ONE TABLE PER ATTRIBUTE)
# ==================================================
with tab_metrics:
    METRICS = [
        "Statistical Parity Difference",
        "Disparate Impact",
        "Average Odds Difference",
        "Equal Opportunity Difference",
        "Error Rate Difference",
    ]

    for attr in protected_attrs:
        st.markdown(f"### {attr}")

        df_attr = inference["results_by_attr"][attr]
        row = df_attr[df_attr["Model"] == results["selected_model"]]

        if row.empty:
            st.warning("Model not found.")
            continue

        row = row.iloc[0]

        st.dataframe(
            [{"Metric": m, "Observed Value": row.get(m, np.nan)} for m in METRICS],
            use_container_width=True
        )

# ==================================================
# TAB 3 — NARRATIVE INPUTS
# ==================================================
with tab_narrative:

    preproc["user_narratives"]["P14"] = st.text_area(
        "AI System Description",
        preproc["user_narratives"]["P14"],
        height=120,
        help=(
            "Provide a concise but complete description of the AI system, "
            "including its purpose, deployment context, users, and decision impact."
        ),
    )

    preproc["pipeline_description"] = st.text_area(
        "Data, Model, and Pipeline Description",
        preproc["pipeline_description"],
        height=120,
        help=(
            "Describe the data sources, preprocessing steps, model type(s), "
            "training procedure, and end-to-end inference pipeline."
        ),
    )

    survey["risk_outcome"] = st.text_area(
        "Risk Assessment Outcome",
        survey["risk_outcome"],
        height=120,
        help=(
            "Summarize the key risks identified during the fairness assessment "
            "and their overall severity."
        ),
    )

    survey["protected_attr_rationale"] = st.text_area(
        "Protected Attribute Rationale",
        survey["protected_attr_rationale"],
        height=120,
        help=(
            "Explain why the selected protected attributes are relevant for this "
            "use case from a legal, ethical, or societal perspective."
        ),
    )

    survey["certification_context"] = st.text_area(
        "Certification Context",
        survey["certification_context"],
        height=120,
        help=(
            "Specify whether this assessment supports regulatory compliance, "
            "internal audit, external certification, or governance reporting."
        ),
    )
# ==================================================
# TAB 4 — REPORT GENERATION
# ==================================================
with tab_report:
    st.markdown("## Generate Final Report")

    if st.button("Generate Final Report"):

        # ------------------------------
        # SELECT WIREFRAME TEMPLATE (BY ATTRIBUTE COUNT)
        # ------------------------------
        attr_count = len(protected_attrs)

        if attr_count == 1:
            template_name = "Fairness_Evaluation_Report_Wireframe_v1.docx"
        elif attr_count == 2:
            template_name = "Fairness_Evaluation_Report_Wireframe_v2.docx"
        elif attr_count == 3:
            template_name = "Fairness_Evaluation_Report_Wireframe_v3.docx"
        else:
            st.error("Unsupported number of protected attributes.")
            st.stop()

        TEMPLATE_PATH = Path(__file__).parent / template_name
        if not TEMPLATE_PATH.exists():
            st.error(f"DOCX wireframe not found: {template_name}")
            st.stop()

        doc = Document(TEMPLATE_PATH)

        # ------------------------------
        # TEXT FIELDS (SUMMARY + HEADER)
        # ------------------------------

        # ------------------------------
        # TEXT FIELDS (SUMMARY + HEADER)
        # ------------------------------
        TEXT = {
            # Header (wireframe top)
            "[[S1]]": survey.get("organization_name", ""),
            "[[S2]]": survey.get("developing_department", ""),

            # Summary section
            "[[P1_TEXT]]": (
                f"System Fairness Score = {results['FS_system']:.3f}. "
            ),
            "[[P2_TEXT]]": survey.get("ai_application_overview", ""),
            "[[P3_TEXT]]": ", ".join(protected_attrs),
            "[[P4_TEXT]]": survey.get("privileged_groups", ""),
            "[[P5_TEXT]]": survey.get("favourable_outcome", ""),

            # Audit / governance
            "[[P6_TEXT]]": survey.get("risk_bucket", ""),
            "[[P7_TEXT]]": survey.get("auditor_access", ""),
            "[[P8_TEXT]]": survey.get("testing_type", ""),
            "[[P9_TEXT]]": survey.get("dependencies", ""),
            "[[P10_TEXT]]": survey.get("limitations", ""),

            # Narrative sections
            "[[P14_TEXT]]": preproc["user_narratives"].get("P14", ""),
            "[[P15_TEXT]]": preproc.get("pipeline_description", ""),
            "[[P16_TEXT]]": build_survey_qa_section_from_submission(survey),
            "[[P17_TEXT]]": survey.get("risk_outcome", ""),
            "[[P18_TEXT]]": survey.get("protected_attr_rationale", ""),
            "[[P26_TEXT]]": survey.get("certification_context", ""),

            # System score
            "[[FS_SYSTEM]]": f"{results['FS_system']:.3f}",
        }

        for k, v in TEXT.items():
            replace_text(doc, k, v)

        # ------------------------------
        # FAIRNESS METRICS + BIAS INDICES
        # ------------------------------
        METRIC_KEYS = {
            "Statistical Parity Difference": "SPD",
            "Disparate Impact": "DI",
            "Average Odds Difference": "AOD",
            "Equal Opportunity Difference": "EOD",
            "Error Rate Difference": "ERD",
        }

        # thresholds selected by user (may be empty)
        thresholds = st.session_state.get("thresholds", {})

        for idx, attr in enumerate(protected_attrs, start=1):

            # Protected attribute name
            replace_text(doc, f"[[ATTR_{idx}_NAME]]", attr)

            # Bias Index
            replace_text(
                doc,
                f"[[BI_ATTR{idx}]]",
                f"{results['BI_per_attribute'][attr]:.3f}"
            )

            df_attr = inference["results_by_attr"][attr]
            row = df_attr[df_attr["Model"] == results["selected_model"]]

            if row.empty:
                continue

            row = row.iloc[0]

            for metric_name, short in METRIC_KEYS.items():

                # ------------------
                # Observed value
                # ------------------
                val = row.get(metric_name, np.nan)
                replace_text(
                    doc,
                    f"[[{short}_ATTR{idx}]]",
                    f"{val:.3f}" if isinstance(val, (float, int)) else "N/A"
                )

                # ------------------
                # Threshold value
                # ------------------
                th_val = "N/A"
                if (
                    attr in thresholds
                    and metric_name in thresholds[attr]
                    and "value" in thresholds[attr][metric_name]
                ):
                    th_val = f"{thresholds[attr][metric_name]['value']:.3f}"

                replace_text(
                    doc,
                    f"[[{short}_TH_ATTR{idx}]]",
                    th_val
                )

        # ------------------------------
        # INSERT PLOTS
        # ------------------------------
        y_true = inference["y_true"]
        y_pred = df[results["selected_model"]].values

        for idx, attr in enumerate(protected_attrs, start=1):
            sens = df[attr].astype(str).values

            fig_disp, _ = plot_disparity_in_performance(
                y_true, y_pred, sens
            )
            insert_plot(doc, f"[[FIG_DISPARITY_ATTR{idx}]]", fig_disp)

            fig_err = plot_group_error_panel(
                y_true, y_pred, sens, group_name=attr
            )
            insert_plot(doc, f"[[FIG_GROUP_ERROR_ATTR{idx}]]", fig_err)

        # ------------------------------
        # SAVE & DOWNLOAD
        # ------------------------------
        out_path = Path(tempfile.gettempdir()) / "Fairness_Evaluation_Report_Final.docx"
        doc.save(out_path)

        st.success("Final fairness evaluation report generated successfully.")
        st.download_button(
            "Download Report",
            open(out_path, "rb"),
            file_name=out_path.name
        )

        # ------------------------------
        # PLOTS (UI preview, unchanged)
        # ------------------------------
        for idx, attr in enumerate(protected_attrs, start=1):
            sens = df[attr].astype(str).values

            fig1, _ = plot_disparity_in_performance(y_true, y_pred, sens)
            insert_plot(doc, f"[[FIG_DISPARITY_ATTR{idx}]]", fig1)

            fig2 = plot_group_error_panel(y_true, y_pred, sens, group_name=attr)
            insert_plot(doc, f"[[FIG_GROUP_ERROR_ATTR{idx}]]", fig2)

        out = Path(tempfile.gettempdir()) / "Fairness_Evaluation_Report_Final.docx"
        doc.save(out)

        st.success("Report generated.")
        st.download_button(
            "Download Report",
            open(out, "rb"),
            file_name=out.name
        )

        # ------------------------------
        # PLOTS
        # ------------------------------
        y_true = inference["y_true"]
        y_pred = df[results["selected_model"]].values

        for idx, attr in enumerate(protected_attrs, start=1):
            sens = df[attr].astype(str).values

            fig1, _ = plot_disparity_in_performance(y_true, y_pred, sens)
            insert_plot(doc, f"[[FIG_DISPARITY_ATTR{idx}]]", fig1)

            fig2 = plot_group_error_panel(y_true, y_pred, sens, group_name=attr)
            insert_plot(doc, f"[[FIG_GROUP_ERROR_ATTR{idx}]]", fig2)

        out = Path(tempfile.gettempdir()) / "Fairness_Evaluation_Report_Final.docx"
        doc.save(out)

        st.success("Report generated.")
        st.download_button("Download Report", open(out, "rb"), file_name=out.name)
