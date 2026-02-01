# pages/5_Report.py
# FINAL REPORT COMPILER — WIRE-FRAME ALIGNED (TEC 7.1 COMPLIANT)


import streamlit as st
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches
import tempfile
import matplotlib.pyplot as plt

from utils.viz_utils import (
    plot_disparity_in_performance,
    plot_group_error_panel,
)



# ==================================================
# PAGE CONFIG
# ==================================================
st.set_page_config(layout="wide")
st.title("Final Fairness Evaluation Report")
st.markdown(
    """
    <style>
    /* ======================================================
       IMPORTS & ROOT VARIABLES
       ====================================================== */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    :root {
        --primary: #2563eb;
        --primary-dark: #1e40af;
        --primary-light: #3b82f6;
        --accent: #0ea5e9;
        --success: #10b981;
        --warning: #f59e0b;
        --danger: #ef4444;
        
        --bg-app: #f8fafc;
        --bg-card: #ffffff;
        --bg-section: #f1f5f9;
        --bg-hover: #e0f2fe;
        
        --border: #e2e8f0;
        --border-strong: #cbd5e1;
        --shadow-sm: 0 1px 3px 0 rgba(0, 0, 0, 0.1);
        --shadow-md: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        --shadow-lg: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
        
        --text-primary: #0f172a;
        --text-secondary: #475569;
        --text-muted: #64748b;
        --text-light: #94a3b8;
    }

    /* ======================================================
       GLOBAL BASE
       ====================================================== */
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif !important;
    }

    .stApp {
        background: linear-gradient(135deg, #f8fafc 0%, #e0f2fe 100%);
        color: var(--text-primary);
    }

    html, body {
        font-size: 16px;
        line-height: 1.6;
        -webkit-font-smoothing: antialiased;
        -moz-osx-font-smoothing: grayscale;
    }

    /* ======================================================
       TYPOGRAPHY
       ====================================================== */
    h1 {
        font-size: 2.5rem !important;
        font-weight: 700 !important;
        letter-spacing: -0.025em !important;
        margin-bottom: 0.5rem !important;
        color: var(--text-primary) !important;
        background: linear-gradient(135deg, #2563eb 0%, #0ea5e9 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }

    h1::after {
        content: "";
        display: block;
        width: 80px;
        height: 4px;
        background: linear-gradient(90deg, var(--primary) 0%, var(--accent) 100%);
        margin-top: 12px;
        border-radius: 2px;
        box-shadow: 0 2px 4px rgba(37, 99, 235, 0.3);
    }

    h2 {
        font-size: 1.75rem !important;
        font-weight: 700 !important;
        margin-top: 2.5rem !important;
        margin-bottom: 1rem !important;
        color: var(--text-primary) !important;
        padding-left: 1rem !important;
        border-left: 5px solid var(--primary) !important;
        background: linear-gradient(90deg, rgba(37, 99, 235, 0.05) 0%, transparent 100%);
        padding: 0.75rem 0 0.75rem 1rem !important;
        border-radius: 0 8px 8px 0;
    }

    h3 {
        font-size: 1.25rem !important;
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        margin-bottom: 0.75rem !important;
    }

    p, .stMarkdown p {
        color: var(--text-secondary) !important;
        font-size: 1rem !important;
        line-height: 1.7 !important;
        max-width: 75ch;
    }

    .stCaption {
        color: var(--text-muted) !important;
        font-size: 0.9rem !important;
    }

    /* ======================================================
       SURVEY NAVIGATION SIDEBAR
       ====================================================== */
    div[data-testid="stVerticalBlockBorderWrapper"]:has(button) {
        background: var(--bg-card) !important;
        border: 2px solid var(--border) !important;
        border-radius: 12px !important;
        padding: 1.5rem !important;
        margin-bottom: 1.5rem !important;
        box-shadow: var(--shadow-md) !important;
    }

    div[data-testid="stVerticalBlockBorderWrapper"] h3 {
        font-size: 1.1rem !important;
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        margin-bottom: 1rem !important;
        padding-bottom: 0.75rem !important;
        border-bottom: 2px solid var(--border) !important;
    }

    /* Navigation buttons */
    div[data-testid="stVerticalBlockBorderWrapper"] button {
        width: 100% !important;
        text-align: left !important;
        padding: 0.75rem 1rem !important;
        margin-bottom: 0.5rem !important;
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
        background: var(--bg-section) !important;
        color: var(--text-secondary) !important;
        font-size: 0.95rem !important;
        font-weight: 500 !important;
        transition: all 0.2s ease !important;
    }

    div[data-testid="stVerticalBlockBorderWrapper"] button:hover {
        background: var(--bg-hover) !important;
        border-color: var(--primary) !important;
        color: var(--primary) !important;
        transform: translateX(4px);
        box-shadow: var(--shadow-sm) !important;
    }

    /* ======================================================
       MAIN CONTENT CARDS
       ====================================================== */
    div[data-testid="stVerticalBlock"] > div[data-testid="stVerticalBlockBorderWrapper"] {
        background: var(--bg-card) !important;
        border: 1px solid var(--border) !important;
        border-radius: 16px !important;
        padding: 2rem !important;
        margin-bottom: 1.5rem !important;
        box-shadow: var(--shadow-md) !important;
        transition: all 0.3s ease;
    }

    div[data-testid="stVerticalBlock"] > div[data-testid="stVerticalBlockBorderWrapper"]:hover {
        box-shadow: var(--shadow-lg) !important;
        border-color: var(--border-strong) !important;
    }

    /* ======================================================
       SURVEY QUESTIONS (EXPANDERS) - KEY IMPROVEMENTS
       ====================================================== */
    div[data-testid="stExpander"] {
        margin-bottom: 1.25rem !important;
        border-radius: 12px !important;
        overflow: hidden;
        box-shadow: var(--shadow-sm) !important;
        transition: all 0.3s ease;
    }

    div[data-testid="stExpander"]:hover {
        box-shadow: var(--shadow-md) !important;
    }

    /* Question headers */
    div[data-testid="stExpander"] > details > summary {
        font-size: 1.05rem !important;
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%) !important;
        padding: 1.25rem 1.5rem !important;
        border-radius: 12px !important;
        border: 2px solid var(--border) !important;
        border-left: 6px solid var(--primary) !important;
        cursor: pointer !important;
        transition: all 0.2s ease !important;
        line-height: 1.6 !important;
        min-height: 60px;
        display: flex;
        align-items: center;
    }

    div[data-testid="stExpander"] > details > summary:hover {
        background: linear-gradient(135deg, var(--bg-hover) 0%, #e0f2fe 100%) !important;
        border-left-color: var(--accent) !important;
        transform: translateX(4px);
        box-shadow: var(--shadow-sm) !important;
    }

    div[data-testid="stExpander"] > details[open] > summary {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%) !important;
        color: white !important;
        border-color: var(--primary-dark) !important;
        border-left-color: var(--accent) !important;
        border-radius: 12px 12px 0 0 !important;
    }

    /* Question content area */
    div[data-testid="stExpander"] > details > div {
        background-color: var(--bg-card) !important;
        padding: 1.5rem 1.75rem !important;
        border-left: 6px solid var(--border-strong) !important;
        border-right: 2px solid var(--border) !important;
        border-bottom: 2px solid var(--border) !important;
        border-radius: 0 0 12px 12px !important;
        box-shadow: inset 0 2px 4px rgba(0, 0, 0, 0.02);
    }

    /* ======================================================
       FORM INPUTS & CONTROLS
       ====================================================== */
    label {
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        font-size: 0.95rem !important;
        margin-bottom: 0.5rem !important;
        display: block !important;
    }

    /* Text inputs and textareas */
    textarea, 
    input[type="text"],
    input[type="number"] {
        font-size: 1rem !important;
        line-height: 1.6 !important;
        border-radius: 8px !important;
        padding: 0.75rem 1rem !important;
        border: 2px solid var(--border) !important;
        background: var(--bg-card) !important;
        color: var(--text-primary) !important;
        transition: all 0.2s ease !important;
        width: 100% !important;
    }

    textarea:focus,
    input[type="text"]:focus,
    input[type="number"]:focus {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.1) !important;
        outline: none !important;
    }

    /* Radio buttons and checkboxes */
    div[data-testid="stRadio"] > div,
    div[data-testid="stCheckbox"] > div {
        background: var(--bg-section) !important;
        padding: 1rem !important;
        border-radius: 8px !important;
        border: 1px solid var(--border) !important;
        margin-bottom: 0.5rem !important;
        transition: all 0.2s ease !important;
    }

    div[data-testid="stRadio"] > div:hover,
    div[data-testid="stCheckbox"] > div:hover {
        background: var(--bg-hover) !important;
        border-color: var(--primary) !important;
    }

    /* Radio button labels */
    div[data-testid="stRadio"] label {
        font-weight: 500 !important;
        font-size: 0.95rem !important;
        padding: 0.5rem 0 !important;
    }

    /* ======================================================
       BUTTONS
       ====================================================== */
    button[kind="primary"] {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%) !important;
        color: white !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        border-radius: 10px !important;
        padding: 0.875rem 2rem !important;
        border: none !important;
        box-shadow: var(--shadow-md) !important;
        transition: all 0.3s ease !important;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }

    button[kind="primary"]:hover {
        background: linear-gradient(135deg, var(--primary-dark) 0%, var(--primary) 100%) !important;
        transform: translateY(-2px);
        box-shadow: var(--shadow-lg) !important;
    }

    button[kind="secondary"] {
        background: var(--bg-card) !important;
        color: var(--text-primary) !important;
        border: 2px solid var(--border) !important;
        font-weight: 600 !important;
        border-radius: 10px !important;
        padding: 0.875rem 2rem !important;
        transition: all 0.2s ease !important;
    }

    button[kind="secondary"]:hover {
        border-color: var(--primary) !important;
        color: var(--primary) !important;
        background: var(--bg-hover) !important;
    }

    /* ======================================================
       METRICS & INFO BOXES
       ====================================================== */
    div[data-testid="stMetric"] {
        background: linear-gradient(135deg, var(--bg-card) 0%, var(--bg-section) 100%) !important;
        padding: 1.5rem !important;
        border-radius: 12px !important;
        border: 2px solid var(--border) !important;
        box-shadow: var(--shadow-md) !important;
    }

    div[data-testid="stMetric"] label {
        font-size: 0.9rem !important;
        font-weight: 600 !important;
        color: var(--text-secondary) !important;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }

    div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
        font-size: 2.5rem !important;
        font-weight: 700 !important;
        color: var(--primary) !important;
    }

    /* Info/Success/Warning boxes */
    div[data-testid="stAlert"] {
        border-radius: 12px !important;
        border: none !important;
        box-shadow: var(--shadow-sm) !important;
        padding: 1.25rem 1.5rem !important;
    }

    div[data-testid="stAlert"][data-baseweb="notification"][kind="info"] {
        background: linear-gradient(135deg, #dbeafe 0%, #e0f2fe 100%) !important;
        border-left: 5px solid var(--accent) !important;
    }

    div[data-testid="stAlert"][data-baseweb="notification"][kind="success"] {
        background: linear-gradient(135deg, #d1fae5 0%, #a7f3d0 100%) !important;
        border-left: 5px solid var(--success) !important;
    }

    div[data-testid="stAlert"][data-baseweb="notification"][kind="warning"] {
        background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%) !important;
        border-left: 5px solid var(--warning) !important;
    }

    div[data-testid="stAlert"][data-baseweb="notification"][kind="error"] {
        background: linear-gradient(135deg, #fee2e2 0%, #fecaca 100%) !important;
        border-left: 5px solid var(--danger) !important;
    }

    /* ======================================================
       SIDEBAR
       ====================================================== */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #ffffff 0%, #f8fafc 100%) !important;
        border-right: 2px solid var(--border) !important;
        box-shadow: 2px 0 8px rgba(0, 0, 0, 0.05);
    }

    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3 {
        color: var(--text-primary) !important;
        padding-left: 0.5rem !important;
    }

    section[data-testid="stSidebar"] li {
        border-radius: 8px !important;
        margin-bottom: 0.25rem !important;
        transition: all 0.2s ease !important;
    }

    section[data-testid="stSidebar"] li:has(a[aria-current="page"]) {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%) !important;
        box-shadow: var(--shadow-sm) !important;
    }

    section[data-testid="stSidebar"] li:has(a[aria-current="page"]) a {
        color: white !important;
        font-weight: 600 !important;
    }

    section[data-testid="stSidebar"] li:hover {
        background: var(--bg-hover) !important;
    }

    /* ======================================================
       PROGRESS & INDICATORS
       ====================================================== */
    .stProgress > div > div {
        background: linear-gradient(90deg, var(--primary) 0%, var(--accent) 100%) !important;
        border-radius: 10px !important;
        height: 8px !important;
    }

    /* ======================================================
       DIVIDERS
       ====================================================== */
    hr {
        border: none !important;
        border-top: 2px solid var(--border) !important;
        margin: 2.5rem 0 !important;
        opacity: 0.6;
    }

    /* ======================================================
       SCROLLBAR STYLING
       ====================================================== */
    ::-webkit-scrollbar {
        width: 10px;
        height: 10px;
    }

    ::-webkit-scrollbar-track {
        background: var(--bg-section);
        border-radius: 10px;
    }

    ::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, var(--primary) 0%, var(--accent) 100%);
        border-radius: 10px;
        border: 2px solid var(--bg-section);
    }

    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(180deg, var(--primary-dark) 0%, var(--primary) 100%);
    }

    /* ======================================================
       RESPONSIVE IMPROVEMENTS
       ====================================================== */
    @media (max-width: 768px) {
        h1 {
            font-size: 2rem !important;
        }
        
        h2 {
            font-size: 1.5rem !important;
        }
        
        div[data-testid="stExpander"] > details > summary {
            font-size: 0.95rem !important;
            padding: 1rem 1.25rem !important;
        }
    }

    /* ======================================================
       ANIMATIONS
       ====================================================== */
    @keyframes slideIn {
        from {
            opacity: 0;
            transform: translateY(10px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    div[data-testid="stVerticalBlock"] > div {
        animation: slideIn 0.3s ease-out;
    }

    /* ======================================================
       FOCUS STATES
       ====================================================== */
    button:focus-visible,
    input:focus-visible,
    textarea:focus-visible {
        outline: 3px solid rgba(37, 99, 235, 0.5) !important;
        outline-offset: 2px !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)



# ==================================================
# IMPORTS REQUIRED FOR THIS PAGE LOGIC
# ==================================================
import numpy as np
import tempfile
import matplotlib.pyplot as plt
from pathlib import Path
from docx import Document
from docx.shared import Inches

from utils.viz_utils import (
    plot_disparity_in_performance,
    plot_group_error_panel,
)

# ==================================================
# DOCX HELPER FUNCTIONS (NO EXTERNAL DEPENDENCIES)
# ==================================================
def replace_text(doc: Document, placeholder: str, value: str) -> bool:
    found = False

    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, str(value))
            found = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, str(value))
                        found = True

    return found


def insert_plot(doc: Document, placeholder: str, fig) -> bool:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(tmp.name, dpi=300, bbox_inches="tight")
    plt.close(fig)

    found = False

    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, "")
            p.add_run().add_picture(tmp.name, width=Inches(6.5))
            found = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        p.add_run().add_picture(tmp.name, width=Inches(6.5))
                        found = True

    return found

def build_survey_qa_section_from_submission(survey_outputs):
    """
    Build Section 3.3 text from stored survey submission.
    Uses section names + question IDs + responses.
    No dependency on survey schema.
    """

    submission = survey_outputs.get("raw_submission", {})
    answers = submission.get("answers", {})

    if not answers:
        return "No survey responses were recorded."

    lines = []
    lines.append("Fairness and Governance Questionnaire Responses:\n")

    for section, qmap in answers.items():
        lines.append(f"{section.replace('_', ' ').title()}:")
        for q_id, response in qmap.items():
            lines.append(f"  - {q_id}: {response}")
        lines.append("")

    return "\n".join(lines)


# ==================================================
# SESSION STATE — SINGLE SOURCE OF TRUTH
# ==================================================
REQUIRED_KEYS = [
    "survey_outputs",
    "preproc",
    "inference",
    "results",
    "uploaded_data",
]

for k in REQUIRED_KEYS:
    if k not in st.session_state:
        st.error(f"Missing required step: {k}")
        st.stop()

survey = st.session_state["survey_outputs"]
preproc = st.session_state["preproc"]
inference = st.session_state["inference"]
results = st.session_state["results"]
df = st.session_state["uploaded_data"]

# ==================================================
# DEFENSIVE INITIALIZATION (ABSOLUTE)
# ==================================================
survey.setdefault("risk_bucket", "Medium")
survey.setdefault("auditor_access", "")
survey.setdefault("testing_type", "")
survey.setdefault("dependencies", "")
survey.setdefault("limitations", "")
survey.setdefault("questionnaire_summary", "")
survey.setdefault("privileged_groups", "")
survey.setdefault("favourable_outcome", "")
survey.setdefault("protected_attr_rationale", "")
survey.setdefault("metric_rationale", "")
survey.setdefault("threshold_rationale", "")
survey.setdefault("risk_outcome", "")
survey.setdefault("certification_context", "")

preproc.setdefault("user_narratives", {})
preproc["user_narratives"].setdefault("P14", "")
preproc.setdefault("pipeline_description", "")

# ==================================================
# HARD RESULTS CONTRACT (NEW FAIRNESS MATH)
# ==================================================
for k in ["selected_model", "FS_system", "BI_per_attribute", "verdict"]:
    if k not in results:
        st.error(f"Results object missing key: {k}")
        st.stop()

protected_attrs = list(results["BI_per_attribute"].keys())
if not protected_attrs:
    st.error("No protected attributes found.")
    st.stop()

# ==================================================
# PAGE TABS
# ==================================================
tab_summary, tab_metrics, tab_narrative, tab_report = st.tabs(
    ["Summary", "Fairness Metrics", "Narrative Inputs", "Generate Report"]
)

# ==================================================
# ==================================================
# TAB 1 — SUMMARY
# ==================================================
with tab_summary:

    # ==================================================
    # ORGANIZATIONAL CONTEXT
    # ==================================================
    st.markdown("## Organizational Context")

    with st.container(border=True):

        survey["organization_name"] = st.text_input(
            "Organization / Entity",
            survey.get("organization_name", ""),
            help=(
                "Legal entity or organization responsible for the development, "
                "deployment, and governance of this AI system."
            ),
        )

        survey["developing_department"] = st.text_input(
            "Department / Business Unit",
            survey.get("developing_department", ""),
            help=(
                "Internal department, team, or business unit that owns or "
                "operates the AI system."
            ),
        )

    # ==================================================
    # AI APPLICATION OVERVIEW
    # ==================================================
    st.markdown("## AI Application Overview")

    with st.container(border=True):

        survey["ai_application_overview"] = st.text_area(
            "High-level description of the AI application",
            survey.get("ai_application_overview", ""),
            height=100,
            help=(
                "Describe the purpose, scope, target users, and decision context "
                "of the AI system."
            ),
        )

    # ==================================================
    # FAIRNESS CONTEXT
    # ==================================================
    st.markdown("## Fairness Context")

    with st.container(border=True):

        survey["privileged_groups"] = st.text_area(
            "Privileged / Reference Groups",
            survey.get("privileged_groups", ""),
            height=80,
            help=(
                "Reference or baseline groups used for fairness comparison "
                "(e.g., majority group, control cohort)."
            ),
        )

        survey["favourable_outcome"] = st.text_area(
            "Definition of Favourable Outcome",
            survey.get("favourable_outcome", ""),
            height=80,
            help=(
                "Define what constitutes a positive or beneficial outcome "
                "in this AI system."
            ),
        )

    # ==================================================
    # AUDIT & GOVERNANCE CONTEXT  
    # ==================================================
    st.markdown("## Audit and Governance Context")

    with st.container(border=True):

        survey["risk_bucket"] = st.selectbox(
            "Audit Type / Risk Classification",
            ["Low", "Medium", "High"],
            index=["Low", "Medium", "High"].index(
                survey.get("risk_bucket", "Medium")
            ),
            help=(
                "Overall audit classification based on regulatory exposure, "
                "deployment criticality, and potential harm."
            ),
        )

        survey["auditor_access"] = st.text_area(
            "Auditor Access to Data",
            survey.get("auditor_access", ""),
            height=80,
            help=(
                "Describe the level of access provided to internal or external "
                "auditors (e.g., full data, anonymized samples, metrics only)."
            ),
        )

        survey["testing_type"] = st.text_area(
            "Testing Type",
            survey.get("testing_type", ""),
            height=80,
            help=(
                "Specify the evaluation methodology used "
                "(e.g., pre-deployment testing, post-deployment monitoring, "
                "periodic audit, shadow testing)."
            ),
        )

    # ==================================================
    # SYSTEM DEPENDENCIES & LIMITATIONS
    # ==================================================
    st.markdown("## System Dependencies and Limitations")

    with st.container(border=True):

        survey["dependencies"] = st.text_area(
            "Dependencies",
            survey.get("dependencies", ""),
            height=80,
            help=(
                "Key technical, data, infrastructure, or organizational "
                "dependencies required for correct operation."
            ),
        )

        survey["limitations"] = st.text_area(
            "Limitations",
            survey.get("limitations", ""),
            height=80,
            help=(
                "Known constraints, assumptions, data gaps, or evaluation "
                "limitations relevant to fairness assessment."
            ),
        )
# ==================================================
# TAB 2 — FAIRNESS METRICS (ONE TABLE PER ATTRIBUTE)
# ==================================================
with tab_metrics:
    METRICS = [
        "Statistical Parity Difference",
        "Disparate Impact",
        "Average Odds Difference",
        "Equal Opportunity Difference",
        "Error Rate Difference",
    ]

    for attr in protected_attrs:
        st.markdown(f"### {attr}")

        df_attr = inference["results_by_attr"][attr]
        row = df_attr[df_attr["Model"] == results["selected_model"]]

        if row.empty:
            st.warning("Model not found.")
            continue

        row = row.iloc[0]

        st.dataframe(
            [{"Metric": m, "Observed Value": row.get(m, np.nan)} for m in METRICS],
            use_container_width=True
        )

# ==================================================
# TAB 3 — NARRATIVE INPUTS
# ==================================================
with tab_narrative:
    preproc["user_narratives"]["P14"] = st.text_area(
        "AI System Description",
        preproc["user_narratives"]["P14"],
        height=120
    )

    preproc["pipeline_description"] = st.text_area(
        "Data, Model, and Pipeline Description",
        preproc["pipeline_description"],
        height=120
    )

  #  survey["questionnaire_summary"] = st.text_area(
  #      "Fairness QuestioSnnaire Summary",
  #      survey["questionnaire_summary"],
  #      height=120
  #  )

    survey["risk_outcome"] = st.text_area(
        "Risk Assessment Outcome",
        survey["risk_outcome"],
        height=120
    )

    survey["protected_attr_rationale"] = st.text_area(
        "Protected Attribute Rationale",
        survey["protected_attr_rationale"],
        height=120
    )

    survey["certification_context"] = st.text_area(
        "Certification Context",
        survey["certification_context"],
        height=120
    )

# ==================================================
# TAB 4 — REPORT GENERATION
# ==================================================
with tab_report:
    st.markdown("## Generate Final Report")

    if st.button("Generate Final Report"):

        # ------------------------------
        # LOAD WIREFRAME TEMPLATE
        # ------------------------------
        TEMPLATE_PATH = Path(__file__).parent / "Fairness_Evaluation_Report_Wireframe_v1.docx"
        if not TEMPLATE_PATH.exists():
            st.error("DOCX wireframe not found.")
            st.stop()

        doc = Document(TEMPLATE_PATH)

        # ------------------------------
        # TEXT FIELDS (SUMMARY + HEADER)
        # ------------------------------
        TEXT = {
            # Header (wireframe top)
            "[[S1]]": survey.get("organization_name", ""),
            "[[S2]]": survey.get("developing_department", ""),

            # Summary section
            "[[P1_TEXT]]": (
                f"System Fairness Score = {results['FS_system']:.3f}. "
                
            ),
            "[[P2_TEXT]]": survey.get("ai_application_overview", ""),
            "[[P3_TEXT]]": ", ".join(protected_attrs),
            "[[P4_TEXT]]": survey.get("privileged_groups", ""),
            "[[P5_TEXT]]": survey.get("favourable_outcome", ""),

            # Audit / governance
            "[[P6_TEXT]]": survey.get("risk_bucket", ""),
            "[[P7_TEXT]]": survey.get("auditor_access", ""),
            "[[P8_TEXT]]": survey.get("testing_type", ""),
            "[[P9_TEXT]]": survey.get("dependencies", ""),
            "[[P10_TEXT]]": survey.get("limitations", ""),

            # Narrative sections
            "[[P14_TEXT]]": preproc["user_narratives"].get("P14", ""),
            "[[P15_TEXT]]": preproc.get("pipeline_description", ""),
            "[[P16_TEXT]]": build_survey_qa_section_from_submission(survey),
            "[[P17_TEXT]]": survey.get("risk_outcome", ""),
            "[[P18_TEXT]]": survey.get("protected_attr_rationale", ""),
            "[[P26_TEXT]]": survey.get("certification_context", ""),

            # System score
            "[[FS_SYSTEM]]": f"{results['FS_system']:.3f}",
        }

        for k, v in TEXT.items():
            replace_text(doc, k, v)
        # ------------------------------
        # FAIRNESS METRICS + BIAS INDICES
        # ------------------------------
        METRIC_KEYS = {
            "Statistical Parity Difference": "SPD",
            "Disparate Impact": "DI",
            "Average Odds Difference": "AOD",
            "Equal Opportunity Difference": "EOD",
            "Error Rate Difference": "ERD",
        }

        for idx, attr in enumerate(protected_attrs, start=1):
            replace_text(doc, f"[[ATTR_{idx}_NAME]]", attr)
            replace_text(
                doc,
                f"[[BI_ATTR{idx}]]",
                f"{results['BI_per_attribute'][attr]:.3f}"
            )

            df_attr = inference["results_by_attr"][attr]
            row = df_attr[df_attr["Model"] == results["selected_model"]]

            if row.empty:
                continue

            row = row.iloc[0]

            for metric_name, short in METRIC_KEYS.items():
                val = row.get(metric_name, np.nan)
                replace_text(
                    doc,
                    f"[[{short}_ATTR{idx}]]",
                    f"{val:.3f}" if isinstance(val, (float, int)) else "N/A"
                )

        # ------------------------------
        # INSERT PLOTS
        # ------------------------------
        y_true = inference["y_true"]
        y_pred = df[results["selected_model"]].values

        for idx, attr in enumerate(protected_attrs, start=1):
            sens = df[attr].astype(str).values

            fig_disp, _ = plot_disparity_in_performance(
                y_true, y_pred, sens
            )
            insert_plot(doc, f"[[FIG_DISPARITY_ATTR{idx}]]", fig_disp)

            fig_err = plot_group_error_panel(
                y_true, y_pred, sens, group_name=attr
            )
            insert_plot(doc, f"[[FIG_GROUP_ERROR_ATTR{idx}]]", fig_err)

        # ------------------------------
        # SAVE & DOWNLOAD
        # ------------------------------
        out_path = Path(tempfile.gettempdir()) / "Fairness_Evaluation_Report_Final.docx"
        doc.save(out_path)

        st.success("Final fairness evaluation report generated successfully.")
        st.download_button(
            "Download Report",
            open(out_path, "rb"),
            file_name=out_path.name
        )
        # ------------------------------
        # PLOTS
        # ------------------------------
        y_true = inference["y_true"]
        y_pred = df[results["selected_model"]].values

        for idx, attr in enumerate(protected_attrs, start=1):
            sens = df[attr].astype(str).values

            fig1, _ = plot_disparity_in_performance(y_true, y_pred, sens)
            insert_plot(doc, f"[[FIG_DISPARITY_ATTR{idx}]]", fig1)

            fig2 = plot_group_error_panel(y_true, y_pred, sens, group_name=attr)
            insert_plot(doc, f"[[FIG_GROUP_ERROR_ATTR{idx}]]", fig2)

        out = Path(tempfile.gettempdir()) / "Fairness_Evaluation_Report_Final.docx"
        doc.save(out)

        st.success("Report generated.")
        st.download_button("Download Report", open(out, "rb"), file_name=out.name)

